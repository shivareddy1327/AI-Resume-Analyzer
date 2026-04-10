"""
Resume Builder utility for Smart Resume AI
"""
import io
from datetime import datetime


class ResumeBuilder:
    def __init__(self):
        pass

    def generate_resume(self, resume_data: dict) -> bytes:
        """Generate a DOCX resume from form data"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            doc = Document()

            # Set page margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.7)
                section.bottom_margin = Inches(0.7)
                section.left_margin = Inches(0.8)
                section.right_margin = Inches(0.8)

            template = resume_data.get("template", "Modern")
            personal_info = resume_data.get("personal_info", {})

            # --- Header / Name ---
            name_para = doc.add_paragraph()
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            name_run = name_para.add_run(personal_info.get("full_name", "Your Name"))
            name_run.bold = True
            name_run.font.size = Pt(22)

            # Contact line
            contact_parts = []
            if personal_info.get("email"):
                contact_parts.append(personal_info["email"])
            if personal_info.get("phone"):
                contact_parts.append(personal_info["phone"])
            if personal_info.get("location"):
                contact_parts.append(personal_info["location"])
            if personal_info.get("linkedin"):
                contact_parts.append(personal_info["linkedin"])
            if personal_info.get("portfolio"):
                contact_parts.append(personal_info["portfolio"])

            if contact_parts:
                contact_para = doc.add_paragraph(" | ".join(contact_parts))
                contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in contact_para.runs:
                    run.font.size = Pt(10)

            self._add_horizontal_line(doc)

            # --- Summary ---
            summary = resume_data.get("summary", "").strip()
            if summary:
                self._add_section_heading(doc, "Professional Summary")
                doc.add_paragraph(summary)

            # --- Experience ---
            experience = resume_data.get("experience", [])
            if experience:
                self._add_section_heading(doc, "Work Experience")
                for exp in experience:
                    if not exp.get("company") and not exp.get("position"):
                        continue
                    # Company / Position line
                    exp_para = doc.add_paragraph()
                    comp_run = exp_para.add_run(exp.get("company", ""))
                    comp_run.bold = True
                    comp_run.font.size = Pt(11)
                    if exp.get("position"):
                        exp_para.add_run(f" — {exp['position']}")

                    # Date line
                    date_parts = []
                    if exp.get("start_date"):
                        date_parts.append(exp["start_date"])
                    if exp.get("end_date"):
                        date_parts.append(exp["end_date"])
                    if date_parts:
                        date_para = doc.add_paragraph(" – ".join(date_parts))
                        date_para.paragraph_format.space_before = Pt(0)
                        for run in date_para.runs:
                            run.italic = True
                            run.font.size = Pt(10)

                    if exp.get("description"):
                        doc.add_paragraph(exp["description"])

                    for resp in exp.get("responsibilities", []):
                        p = doc.add_paragraph(style='List Bullet')
                        p.add_run(resp)

                    for ach in exp.get("achievements", []):
                        p = doc.add_paragraph(style='List Bullet')
                        p.add_run(f"✓ {ach}")

            # --- Education ---
            education = resume_data.get("education", [])
            if education:
                self._add_section_heading(doc, "Education")
                for edu in education:
                    if not edu.get("school"):
                        continue
                    edu_para = doc.add_paragraph()
                    school_run = edu_para.add_run(edu.get("school", ""))
                    school_run.bold = True
                    school_run.font.size = Pt(11)

                    degree_parts = []
                    if edu.get("degree"):
                        degree_parts.append(edu["degree"])
                    if edu.get("field"):
                        degree_parts.append(edu["field"])
                    if degree_parts:
                        doc.add_paragraph(", ".join(degree_parts))

                    info_parts = []
                    if edu.get("graduation_date"):
                        info_parts.append(f"Graduated: {edu['graduation_date']}")
                    if edu.get("gpa"):
                        info_parts.append(f"GPA: {edu['gpa']}")
                    if info_parts:
                        info_para = doc.add_paragraph(" | ".join(info_parts))
                        for run in info_para.runs:
                            run.italic = True
                            run.font.size = Pt(10)

                    for ach in edu.get("achievements", []):
                        p = doc.add_paragraph(style='List Bullet')
                        p.add_run(ach)

            # --- Projects ---
            projects = resume_data.get("projects", [])
            if projects:
                self._add_section_heading(doc, "Projects")
                for proj in projects:
                    if not proj.get("name"):
                        continue
                    proj_para = doc.add_paragraph()
                    proj_run = proj_para.add_run(proj.get("name", ""))
                    proj_run.bold = True
                    proj_run.font.size = Pt(11)

                    if proj.get("technologies"):
                        proj_para.add_run(f" | {proj['technologies']}")

                    if proj.get("link"):
                        doc.add_paragraph(proj["link"])

                    if proj.get("description"):
                        doc.add_paragraph(proj["description"])

                    for resp in proj.get("responsibilities", []):
                        p = doc.add_paragraph(style='List Bullet')
                        p.add_run(resp)

            # --- Skills ---
            skills = resume_data.get("skills", {})
            if any(skills.get(k) for k in ['technical', 'soft', 'languages', 'tools']):
                self._add_section_heading(doc, "Skills")

                if skills.get("technical"):
                    p = doc.add_paragraph()
                    p.add_run("Technical: ").bold = True
                    p.add_run(", ".join(skills["technical"]))

                if skills.get("tools"):
                    p = doc.add_paragraph()
                    p.add_run("Tools & Technologies: ").bold = True
                    p.add_run(", ".join(skills["tools"]))

                if skills.get("soft"):
                    p = doc.add_paragraph()
                    p.add_run("Soft Skills: ").bold = True
                    p.add_run(", ".join(skills["soft"]))

                if skills.get("languages"):
                    p = doc.add_paragraph()
                    p.add_run("Languages: ").bold = True
                    p.add_run(", ".join(skills["languages"]))

            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.read()

        except Exception as e:
            print(f"Error generating resume: {e}")
            return None

    def _add_section_heading(self, doc, title: str):
        """Add a styled section heading"""
        from docx.shared import Pt, RGBColor
        heading = doc.add_heading(title, level=2)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(4)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    def _add_horizontal_line(self, doc):
        """Add a horizontal line"""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        p = doc.add_paragraph()
        p.paragraph_format.space_before = 0
        p.paragraph_format.space_after = 0
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '1a1a2e')
        pBdr.append(bottom)
        pPr.append(pBdr)
